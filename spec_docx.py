"""Specification Word Doc builder.

Renders a spec_engine.RenderedSpec to a formatted .docx using python-docx.
Calibri font, black and white, matching the on-page preview style.
Notes are stripped (build_spec called with include_notes=False).

Add to requirements.txt:
    python-docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_bottom_border(paragraph):
    """Light grey bottom rule under a paragraph (section heading)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_font(run, size_pt: float, bold: bool = False):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def build_specification_docx(rendered_spec, out_path: Path,
                              project_name: str,
                              project_address: str,
                              code_label: str) -> Path:
    """Write a Calibri-styled .docx spec to out_path and return out_path."""

    doc = Document()

    # ── Page margins: 1 inch all around ──────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Default paragraph spacing: no extra space after ───────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)

    # ── Title block ───────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    _set_font(title_p.add_run("MECHANICAL SPECIFICATIONS"), 14, bold=True)

    if code_label:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(2)
        _set_font(sub_p.add_run(f"{code_label}  \u00b7  Division 23"), 10)

    proj_line = project_name
    if project_address:
        proj_line += f"  \u2014  {project_address}"
    addr_p = doc.add_paragraph()
    addr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr_p.paragraph_format.space_after = Pt(14)
    _set_font(addr_p.add_run(proj_line), 10)

    # ── Spec content ──────────────────────────────────────────────────
    for part in rendered_spec.parts:
        for section in part.sections:

            # Section heading — bold, uppercase, bottom border
            sec_p = doc.add_paragraph()
            sec_p.paragraph_format.space_before = Pt(10)
            sec_p.paragraph_format.space_after  = Pt(3)
            _set_font(sec_p.add_run(
                f"{section.num}  {section.title.upper()}"
            ), 11, bold=True)
            _add_bottom_border(sec_p)

            # Clauses — hanging indent, label bold, text normal
            for clause in section.clauses:
                c_p = doc.add_paragraph(style="Normal")
                c_p.paragraph_format.left_indent       = Inches(0.4)
                c_p.paragraph_format.first_line_indent = Inches(-0.4)
                c_p.paragraph_format.space_after        = Pt(3)

                label_run = c_p.add_run(f"{clause.label}  ")
                _set_font(label_run, 11, bold=True)

                text_run = c_p.add_run(clause.text)
                _set_font(text_run, 11)

            # Blank line between sections
            gap = doc.add_paragraph()
            gap.paragraph_format.space_after = Pt(0)

    doc.save(str(out_path))
    return out_path
